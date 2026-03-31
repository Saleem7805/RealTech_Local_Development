from flask import Flask, request, jsonify
import os
import fitz  # PyMuPDF
import re
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
import docx


# ---------------------- CONFIG ----------------------
SECTION_PATTERNS = {
    "education": [r"\beducation\b", r"\bqualification\b"],
    "experience": [r"\bexperience\b", r"\bwork history\b"],
    "skills": [r"\bskills\b", r"\btechnical skills\b"],
    "projects": [r"\bprojects\b"],
    "secondary_skills": [r"\badditional skills\b", r"\bother skills\b"],
    "domain": [r"\bdomain\b", r"\bindustry\b"]
}

# ---------------------- LOGGING ----------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------------- NLP MODEL ----------------------
try:
    nlp = spacy.load("en_core_web_sm")
except:
    raise RuntimeError("Run: python -m spacy download en_core_web_sm")

# ---------------------- FLASK APP ----------------------
app = Flask(__name__)

# ---------------------- PDF PARSER ----------------------
# def extract_text(file_path: str) -> str:
#     try:
#         if file_path.lower().endswith(".pdf"):
#             import fitz
#             doc = fitz.open(file_path)
#             text = ""
#             for page in doc:
#                 text += page.get_text()
#             return text

#         elif file_path.lower().endswith(".docx"):
#             doc = docx.Document(file_path)
#             return "\n".join([para.text for para in doc.paragraphs])

        
#         else:
#             return ""

#     except Exception as e:
#         logger.error(f"File read error: {file_path} | {e}")
#         return ""



import docx

def extract_text(file_path: str) -> str:
    try:
        if file_path.lower().endswith(".docx"):
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        
        elif file_path.lower().endswith(".doc"):
            logger.warning(f"Skipping .doc file (convert to .docx): {file_path}")
            return ""

        else:
            return ""

    except Exception as e:
        logger.error(f"File read error: {file_path} | {e}")
        return ""
# ---------------------- SECTION DETECTOR ----------------------
def detect_sections(text: str) -> dict:
    sections = {}
    current_section = "other"
    sections[current_section] = []

    lines = text.split("\n")

    for line in lines:
        line_clean = line.strip().lower()
        matched = False

        for section, patterns in SECTION_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, line_clean):
                    current_section = section
                    sections.setdefault(section, [])
                    matched = True
                    break
            if matched:
                break

        sections.setdefault(current_section, []).append(line)

    return {k: " ".join(v).strip() for k, v in sections.items()}


# ---------------------- NLP ENTITY EXTRACTION ----------------------
def extract_entities(text: str):
    try:
        doc = nlp(text)
        entities = {"ORG": [], "DATE": []}

        for ent in doc.ents:
            if ent.label_ in ["ORG", "DATE"]:
                entities[ent.label_].append(ent.text)

        return entities
    except:
        return {"ORG": [], "DATE": []}


# ---------------------- TF-IDF ----------------------
def tfidf_classify(sections: dict):
    try:
        corpus = list(sections.values())
        if not corpus:
            return []

        vectorizer = TfidfVectorizer(stop_words="english")
        tfidf_matrix = vectorizer.fit_transform(corpus)

        return tfidf_matrix.toarray()
    except:
        return []


# ---------------------- MAIN PIPELINE ----------------------
def process_resume(file_path: str) -> dict:
    logger.info(f"Processing file: {file_path}")

    text = extract_text(file_path)
    sections = detect_sections(text)

    return {
        "education": sections.get("education", ""),
        "skill_set": sections.get("skills", ""),
        "experience": sections.get("experience", ""),
        "projects": sections.get("projects", ""),
        "secondary_skill": sections.get("secondary_skills", ""),
        "domain": sections.get("domain", "")
    }


# ---------------------- SINGLE FILE HANDLER ----------------------
def handle_single_file(file_path):
    file_name = os.path.basename(file_path)

    print(f"Starting with {file_name}")

    result = process_resume(file_path)

    print(f"Ending with {file_name}")

    return {
        "file_name": file_name,
        "extracted_fields": result
    }


# ---------------------- API ROUTE ----------------------
@app.route("/extract_section", methods=["POST"])
def extract_sections():
    try:
        data = request.get_json()
        folder_path = data.get("folder_path")

        if not folder_path:
            return jsonify({"error": "folder_path is required"}), 400

        if not os.path.exists(folder_path):
            return jsonify({"error": "Folder not found"}), 400

        doc_files = [
            os.path.join(folder_path, f)
            for f in os.listdir(folder_path)
            if f.lower().endswith((".docx", ".doc"))
        ]

        if not doc_files:
            return jsonify({"error": "No Doc files found in folder"}), 400

        results = []

        # ✅ Simple loop (no parallel)
        for file in doc_files:
            res = handle_single_file(file)
            if res:
                results.append(res)

        return jsonify(results)

    except Exception as e:
        logger.error(f"Error: {e}")
        return jsonify({"error": str(e)}), 500

    finally:
        # ✅ This always runs (success or error)
        logger.info("Request processing completed")
# ---------------------- RUN SERVER ----------------------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)